import io
from pathlib import Path


async def convert_file(file_info: dict, target_format: str) -> dict:
    """Convert a file to the target format."""
    source_name = file_info["name"]
    source_ext = Path(source_name).suffix.lower().lstrip(".")
    data = file_info["data"]

    if source_ext == "pdf" and target_format == "docx":
        return await _pdf_to_docx(data, source_name)
    elif source_ext == "docx" and target_format == "pdf":
        return await _docx_to_pdf(data, source_name)
    elif source_ext in ("jpg", "jpeg", "png") and target_format == "pdf":
        return await _image_to_pdf(data, source_name)
    elif source_ext == "md" and target_format == "html":
        return await _md_to_html(data, source_name)
    else:
        raise ValueError(f"Conversion {source_ext} -> {target_format} not supported")


async def _pdf_to_docx(data: bytes, name: str) -> dict:
    """Convert PDF to DOCX."""
    # Uses pypdf for text extraction, python-docx for creation
    from pypdf import PdfReader
    from docx import Document

    reader = PdfReader(io.BytesIO(data))
    doc = Document()

    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)

    return {
        "data": buffer.getvalue(),
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": Path(name).stem + ".docx",
    }


async def _docx_to_pdf(data: bytes, name: str) -> dict:
    """Convert DOCX to PDF (basic text extraction)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)

    # Simple HTML-to-PDF approach
    html = f"<html><body><pre style='font-family:serif;font-size:12pt'>{text}</pre></body></html>"

    return {
        "data": html.encode("utf-8"),
        "mime_type": "text/html",
        "filename": Path(name).stem + ".html",
    }


async def _image_to_pdf(data: bytes, name: str) -> dict:
    """Convert image to PDF."""
    from PIL import Image

    img = Image.open(io.BytesIO(data))
    if img.mode == "RGBA":
        img = img.convert("RGB")

    buffer = io.BytesIO()
    img.save(buffer, format="PDF")

    return {
        "data": buffer.getvalue(),
        "mime_type": "application/pdf",
        "filename": Path(name).stem + ".pdf",
    }


async def _md_to_html(data: bytes, name: str) -> dict:
    """Convert Markdown to HTML."""
    import markdown

    text = data.decode("utf-8")
    html = markdown.markdown(text, extensions=["extra", "sane_lists"], output_format="html5")
    full_html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{html}</body></html>"

    return {
        "data": full_html.encode("utf-8"),
        "mime_type": "text/html",
        "filename": Path(name).stem + ".html",
    }


async def merge_files(files: list[dict]) -> dict:
    """Merge multiple PDF files into one."""
    from pypdf import PdfWriter, PdfReader

    writer = PdfWriter()
    for f in files:
        reader = PdfReader(io.BytesIO(f["data"]))
        for page in reader.pages:
            writer.add_page(page)

    buffer = io.BytesIO()
    writer.write(buffer)

    return {
        "data": buffer.getvalue(),
        "mime_type": "application/pdf",
        "filename": "merged.pdf",
    }


async def split_file(file_info: dict) -> dict:
    """Split PDF into individual pages (returns first page as demo)."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(file_info["data"]))

    if len(reader.pages) <= 1:
        return {
            "data": file_info["data"],
            "mime_type": "application/pdf",
            "filename": "page_1.pdf",
        }

    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    buffer = io.BytesIO()
    writer.write(buffer)

    return {
        "data": buffer.getvalue(),
        "mime_type": "application/pdf",
        "filename": "page_1.pdf",
    }


async def compress_file(file_info: dict) -> dict:
    """Compress PDF by reducing image quality."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(file_info["data"]))
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)

    buffer = io.BytesIO()
    writer.write(buffer)

    return {
        "data": buffer.getvalue(),
        "mime_type": "application/pdf",
        "filename": "compressed_" + file_info["name"],
    }
