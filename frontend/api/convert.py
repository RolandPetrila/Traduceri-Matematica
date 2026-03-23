"""Vercel Python Serverless Function for File Conversion.

Handles POST /api/convert — convert/merge/split/compress files.
Dependencies: pypdf, python-docx, Pillow, markdown (in requirements.txt).
"""

from __future__ import annotations

from http.server import BaseHTTPRequestHandler
import io
import json
import re
import sys
import traceback


# --- Helpers ---

def _stem(name: str) -> str:
    return name.rsplit(".", 1)[0] if "." in name else name


# --- Conversion functions ---

def pdf_to_docx(data: bytes, name: str) -> dict:
    from pypdf import PdfReader
    from docx import Document

    reader = PdfReader(io.BytesIO(data))
    doc = Document()
    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.add_page_break()
    buf = io.BytesIO()
    doc.save(buf)
    return {
        "data": buf.getvalue(),
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": f"{_stem(name)}.docx",
    }


def pdf_to_html(data: bytes, name: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = (page.extract_text() or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        pages.append(f"<div class='page'><pre>{text}</pre></div>")
    body = "\n<hr>\n".join(pages)
    s = _stem(name)
    html = (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title>"
        "<style>.page{margin:2em 0}pre{font-family:serif;font-size:12pt;white-space:pre-wrap}</style>"
        f"</head><body>{body}</body></html>"
    )
    return {"data": html.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def docx_to_html(data: bytes, name: str) -> dict:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        text = p.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if p.style and p.style.name and "Heading" in p.style.name:
            lvl = p.style.name[-1] if p.style.name[-1].isdigit() else "2"
            parts.append(f"<h{lvl}>{text}</h{lvl}>")
        else:
            parts.append(f"<p>{text}</p>")
    body = "\n".join(parts)
    s = _stem(name)
    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title></head><body>{body}</body></html>"
    return {"data": html.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def image_to_pdf(data: bytes, name: str) -> dict:
    from PIL import Image

    img = Image.open(io.BytesIO(data))
    if img.mode == "RGBA":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PDF")
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"{_stem(name)}.pdf"}


def image_convert(data: bytes, name: str, target: str) -> dict:
    from PIL import Image

    fmt_map = {"jpg": ("JPEG", "image/jpeg"), "jpeg": ("JPEG", "image/jpeg"), "png": ("PNG", "image/png")}
    pil_fmt, mime = fmt_map.get(target, ("PNG", "image/png"))
    img = Image.open(io.BytesIO(data))
    if pil_fmt == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format=pil_fmt)
    return {"data": buf.getvalue(), "mime": mime, "filename": f"{_stem(name)}.{target}"}


def md_to_html(data: bytes, name: str) -> dict:
    import markdown as md_lib

    text = data.decode("utf-8")
    html = md_lib.markdown(text, extensions=["extra", "sane_lists"], output_format="html5")
    s = _stem(name)
    full = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title></head><body>{html}</body></html>"
    return {"data": full.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def merge_pdfs(files: list[dict]) -> dict:
    from pypdf import PdfWriter, PdfReader

    writer = PdfWriter()
    for f in files:
        reader = PdfReader(io.BytesIO(f["data"]))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": "merged.pdf"}


def split_pdf(data: bytes) -> dict:
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) <= 1:
        return {"data": data, "mime": "application/pdf", "filename": "page_1.pdf"}
    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": "page_1.pdf"}


def compress_pdf(data: bytes, name: str) -> dict:
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(data))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)
    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"compressed_{name}"}


# --- Router ---

def process(files: list[dict], operation: str, target_format: str) -> dict:
    if not files:
        raise ValueError("Nu au fost trimise fisiere")

    f0 = files[0]
    name = f0.get("name", "file.bin")
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    data = f0["data"]

    if operation == "merge":
        return merge_pdfs(files)
    if operation == "split":
        return split_pdf(data)
    if operation == "compress":
        return compress_pdf(data, name)

    if operation != "convert" or not target_format:
        raise ValueError(f"Operatiune invalida: {operation}")

    # Image conversions
    if ext in ("jpg", "jpeg", "png") and target_format == "pdf":
        return image_to_pdf(data, name)
    if ext in ("jpg", "jpeg", "png") and target_format in ("jpg", "jpeg", "png"):
        return image_convert(data, name, target_format)

    routes = {
        ("pdf", "docx"): lambda: pdf_to_docx(data, name),
        ("pdf", "html"): lambda: pdf_to_html(data, name),
        ("docx", "html"): lambda: docx_to_html(data, name),
        ("docx", "pdf"): lambda: docx_to_html(data, name),
        ("md", "html"): lambda: md_to_html(data, name),
        ("md", "pdf"): lambda: md_to_html(data, name),
    }
    key = (ext, target_format)
    if key in routes:
        return routes[key]()
    raise ValueError(f"Conversie {ext} -> {target_format} nu e suportata")


# --- Multipart parser ---

def parse_boundary(content_type: str) -> str:
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            return part[len("boundary="):].strip('"').strip("'")
    raise ValueError("No boundary in Content-Type")


def parse_multipart(body: bytes, boundary: str) -> dict:
    result: dict = {"files": [], "operation": "convert", "target_format": ""}
    boundary_bytes = f"--{boundary}".encode()
    sections = body.split(boundary_bytes)

    for section in sections[1:]:
        if section.strip() in (b"--", b"", b"--\r\n"):
            continue
        header_end = section.find(b"\r\n\r\n")
        if header_end == -1:
            continue
        header = section[:header_end].decode("utf-8", errors="replace")
        content = section[header_end + 4:]
        if content.endswith(b"\r\n"):
            content = content[:-2]

        name_match = re.search(r'name="([^"]+)"', header)
        if not name_match:
            continue
        name = name_match.group(1)

        if name in ("operation", "target_format"):
            result[name] = content.decode("utf-8").strip()
        elif name == "files" or "filename" in header:
            fname_match = re.search(r'filename="([^"]*)"', header)
            fname = fname_match.group(1) if fname_match else "file.bin"
            ct_match = re.search(r"Content-Type:\s*(\S+)", header)
            mime = ct_match.group(1) if ct_match else "application/octet-stream"
            result["files"].append({"name": fname, "data": content, "mime": mime})

    return result


# --- Handler ---

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        try:
            content_type = self.headers.get("Content-Type", "")
            content_length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(content_length)

            if "multipart/form-data" in content_type:
                boundary = parse_boundary(content_type)
                parts = parse_multipart(body, boundary)
            else:
                parts = json.loads(body)

            files = parts.get("files", [])
            operation = parts.get("operation", "convert")
            target_format = parts.get("target_format", "")

            print(f"[CONVERT] {operation}: {len(files)} files, target={target_format}", file=sys.stderr)

            result = process(files, operation, target_format)

            self.send_response(200)
            self.send_header("Content-Type", result["mime"])
            self.send_header("Content-Disposition", f'attachment; filename="{result["filename"]}"')
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(result["data"])

        except Exception as e:
            print(f"[CONVERT ERROR] {type(e).__name__}: {e}", file=sys.stderr)
            traceback.print_exc(file=sys.stderr)
            error_body = json.dumps({"error": str(e), "status": "error"}).encode()
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(error_body)
