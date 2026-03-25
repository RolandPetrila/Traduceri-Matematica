"""Vercel Python Serverless Function for File Conversion.

Handles POST /api/convert — convert/merge/split/compress files.
Dependencies: pypdf, python-docx, Pillow, markdown (in requirements.txt).
"""

from __future__ import annotations

from http.server import BaseHTTPRequestHandler
import io
import json
import re
import sys
import traceback



# --- Local debug logging ---


def _log_to_file(message: str) -> None:
    """Append log entry to data/logs/local_debug.log (local dev only)."""
    import os

    log_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "logs"
    )
    if not os.path.isdir(log_dir):
        return
    log_file = os.path.join(log_dir, "local_debug.log")
    try:
        from datetime import datetime

        ts = datetime.now().strftime("%H:%M:%S")
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] {message}\n")
    except Exception:
        pass


# --- Helpers ---

def _stem(name: str) -> str:
    return name.rsplit(".", 1)[0] if "." in name else name


# --- Conversion functions ---

def pdf_to_docx(data: bytes, name: str) -> dict:
    from pypdf import PdfReader
    from docx import Document

    reader = PdfReader(io.BytesIO(data))
    doc = Document()
    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.add_page_break()
    buf = io.BytesIO()
    doc.save(buf)
    return {
        "data": buf.getvalue(),
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": f"{_stem(name)}.docx",
    }


def pdf_to_html(data: bytes, name: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = (page.extract_text() or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        pages.append(f"<div class='page'><pre>{text}</pre></div>")
    body = "\n<hr>\n".join(pages)
    s = _stem(name)
    html = (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title>"
        "<style>.page{margin:2em 0}pre{font-family:serif;font-size:12pt;white-space:pre-wrap}</style>"
        f"</head><body>{body}</body></html>"
    )
    return {"data": html.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def docx_to_html(data: bytes, name: str) -> dict:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        text = p.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if p.style and p.style.name and "Heading" in p.style.name:
            lvl = p.style.name[-1] if p.style.name[-1].isdigit() else "2"
            parts.append(f"<h{lvl}>{text}</h{lvl}>")
        else:
            parts.append(f"<p>{text}</p>")
    body = "\n".join(parts)
    s = _stem(name)
    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title></head><body>{body}</body></html>"
    return {"data": html.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def image_to_pdf(data: bytes, name: str) -> dict:
    from PIL import Image

    img = Image.open(io.BytesIO(data))
    if img.mode == "RGBA":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PDF")
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"{_stem(name)}.pdf"}


def image_convert(data: bytes, name: str, target: str) -> dict:
    from PIL import Image

    fmt_map = {"jpg": ("JPEG", "image/jpeg"), "jpeg": ("JPEG", "image/jpeg"), "png": ("PNG", "image/png")}
    pil_fmt, mime = fmt_map.get(target, ("PNG", "image/png"))
    img = Image.open(io.BytesIO(data))
    if pil_fmt == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format=pil_fmt)
    return {"data": buf.getvalue(), "mime": mime, "filename": f"{_stem(name)}.{target}"}


def md_to_html(data: bytes, name: str) -> dict:
    import markdown as md_lib

    text = data.decode("utf-8")
    html = md_lib.markdown(text, extensions=["extra", "sane_lists"], output_format="html5")
    s = _stem(name)
    full = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{s}</title></head><body>{html}</body></html>"
    return {"data": full.encode("utf-8"), "mime": "text/html", "filename": f"{s}.html"}


def _text_to_pdf(text: str, title: str = "") -> bytes:
    """Convert plain/structured text to PDF using fpdf2 (pure Python)."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    # Use built-in Helvetica (supports latin-1; for full Unicode would need font file)
    pdf.set_font("Helvetica", size=11)
    if title:
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(4)
        pdf.set_font("Helvetica", size=11)
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, stripped[2:], ln=True)
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 9, stripped[3:], ln=True)
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, stripped[4:], ln=True)
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("**") and stripped.endswith("**"):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, stripped[2:-2])
            pdf.set_font("Helvetica", size=11)
        elif not stripped:
            pdf.ln(3)
        else:
            pdf.multi_cell(0, 6, stripped)
    return pdf.output()


def docx_to_pdf(data: bytes, name: str) -> dict:
    """Convert DOCX to PDF by extracting text and rendering with fpdf2."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = (p.style.name if p.style else "").lower()
        if "heading 1" in style_name or "title" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif not text:
            lines.append("")
        else:
            lines.append(text)
    content = "\n".join(lines)
    pdf_bytes = _text_to_pdf(content, _stem(name))
    return {"data": pdf_bytes, "mime": "application/pdf", "filename": f"{_stem(name)}.pdf"}


def md_to_pdf(data: bytes, name: str) -> dict:
    """Convert Markdown to PDF by rendering text with fpdf2."""
    text = data.decode("utf-8")
    pdf_bytes = _text_to_pdf(text, _stem(name))
    return {"data": pdf_bytes, "mime": "application/pdf", "filename": f"{_stem(name)}.pdf"}


def html_to_pdf(data: bytes, name: str) -> dict:
    """Convert HTML to PDF by stripping tags and rendering text."""
    text = data.decode("utf-8")
    # Strip HTML tags, keep text
    clean = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", text)
    clean = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", clean)
    clean = re.sub(r"<br\s*/?>", "\n", clean, flags=re.IGNORECASE)
    clean = re.sub(r"</(p|div|h[1-6]|li|tr)>", "\n", clean, flags=re.IGNORECASE)
    clean = re.sub(r"<[^>]+>", "", clean)
    clean = re.sub(r"&amp;", "&", clean)
    clean = re.sub(r"&lt;", "<", clean)
    clean = re.sub(r"&gt;", ">", clean)
    clean = re.sub(r"&nbsp;", " ", clean)
    clean = re.sub(r"\n{3,}", "\n\n", clean)
    pdf_bytes = _text_to_pdf(clean.strip(), _stem(name))
    return {"data": pdf_bytes, "mime": "application/pdf", "filename": f"{_stem(name)}.pdf"}


def html_to_md(data: bytes, name: str) -> dict:
    """Convert HTML to Markdown with basic structure preservation."""
    text = data.decode("utf-8")
    # Remove script/style
    md = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", text)
    md = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", md)
    # Convert headings
    for i in range(1, 7):
        md = re.sub(rf"<h{i}[^>]*>(.*?)</h{i}>", rf"{'#' * i} \1\n", md, flags=re.IGNORECASE | re.DOTALL)
    # Convert formatting
    md = re.sub(r"<strong>(.*?)</strong>", r"**\1**", md, flags=re.DOTALL)
    md = re.sub(r"<b>(.*?)</b>", r"**\1**", md, flags=re.DOTALL)
    md = re.sub(r"<em>(.*?)</em>", r"*\1*", md, flags=re.DOTALL)
    md = re.sub(r"<i>(.*?)</i>", r"*\1*", md, flags=re.DOTALL)
    # Convert lists
    md = re.sub(r"<li>(.*?)</li>", r"- \1\n", md, flags=re.DOTALL)
    # Line breaks
    md = re.sub(r"<br\s*/?>", "\n", md, flags=re.IGNORECASE)
    md = re.sub(r"</(p|div|tr)>", "\n\n", md, flags=re.IGNORECASE)
    # Strip remaining tags
    md = re.sub(r"<[^>]+>", "", md)
    # Clean entities
    md = re.sub(r"&amp;", "&", md)
    md = re.sub(r"&lt;", "<", md)
    md = re.sub(r"&gt;", ">", md)
    md = re.sub(r"&nbsp;", " ", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = md.strip() + "\n"
    return {"data": md.encode("utf-8"), "mime": "text/markdown", "filename": f"{_stem(name)}.md"}


def html_to_docx(data: bytes, name: str) -> dict:
    """Convert HTML to DOCX with basic structure preservation."""
    from docx import Document
    from docx.shared import Pt

    text = data.decode("utf-8")
    # Strip script/style
    clean = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", text)
    clean = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", clean)

    doc = Document()
    # Extract headings and paragraphs
    blocks = re.split(r"</(p|div|h[1-6]|li|tr)>", clean, flags=re.IGNORECASE)
    for block in blocks:
        # Check for heading
        h_match = re.search(r"<h([1-6])[^>]*>(.*)", block, re.IGNORECASE | re.DOTALL)
        if h_match:
            level = int(h_match.group(1))
            content = re.sub(r"<[^>]+>", "", h_match.group(2)).strip()
            if content:
                doc.add_heading(content, level=min(level, 4))
            continue
        # Strip tags for regular text
        content = re.sub(r"<br\s*/?>", "\n", block, flags=re.IGNORECASE)
        content = re.sub(r"<[^>]+>", "", content)
        content = re.sub(r"&amp;", "&", content)
        content = re.sub(r"&lt;", "<", content)
        content = re.sub(r"&gt;", ">", content)
        content = re.sub(r"&nbsp;", " ", content)
        content = content.strip()
        if content:
            p = doc.add_paragraph(content)
            p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return {
        "data": buf.getvalue(),
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": f"{_stem(name)}.docx",
    }


def merge_pdfs(files: list[dict]) -> dict:
    from pypdf import PdfWriter, PdfReader

    writer = PdfWriter()
    for f in files:
        reader = PdfReader(io.BytesIO(f["data"]))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": "merged.pdf"}


def _parse_page_range(spec: str, total: int) -> list[int]:
    """Parse page range spec like '1-3,5,7-10' into 0-based page indices."""
    if not spec or spec.strip().lower() == "all":
        return list(range(total))
    pages: list[int] = []
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            s = max(1, int(start.strip()))
            e = min(total, int(end.strip()))
            pages.extend(range(s - 1, e))
        else:
            p = int(part) - 1
            if 0 <= p < total:
                pages.append(p)
    return sorted(set(pages)) if pages else list(range(total))


def split_pdf(data: bytes, page_range: str = "") -> dict:
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(data))
    indices = _parse_page_range(page_range, len(reader.pages))
    writer = PdfWriter()
    for idx in indices:
        writer.add_page(reader.pages[idx])
    buf = io.BytesIO()
    writer.write(buf)
    label = page_range.strip() if page_range.strip() else "all"
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"pages_{label}.pdf"}


def compress_pdf(data: bytes, name: str) -> dict:
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(data))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)
    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"compressed_{name}"}


# --- PDF Edit operations ---

def edit_pdf(data: bytes, name: str, **kwargs) -> dict:
    """Handle PDF edit operations: rotate, delete, reorder, optimize, watermark."""
    from pypdf import PdfReader, PdfWriter

    action = kwargs.get("pdf_action", "")
    if not action:
        raise ValueError("Specifica o operatie PDF (rotate/delete/reorder/optimize/watermark)")

    reader = PdfReader(io.BytesIO(data))
    total = len(reader.pages)
    writer = PdfWriter()

    if action == "rotate":
        angle = int(kwargs.get("rotate_angle", "90"))
        indices = _parse_page_range(kwargs.get("page_range", "all"), total)
        for i in range(total):
            page = reader.pages[i]
            if i in indices:
                page.rotate(angle)
            writer.add_page(page)

    elif action == "delete":
        indices = _parse_page_range(kwargs.get("page_range", ""), total)
        if not indices:
            raise ValueError("Specifica paginile de sters")
        keep = [i for i in range(total) if i not in indices]
        if not keep:
            raise ValueError("Nu poti sterge toate paginile")
        for i in keep:
            writer.add_page(reader.pages[i])

    elif action == "reorder":
        seq_str = kwargs.get("reorder_sequence", "")
        if not seq_str:
            raise ValueError("Specifica ordinea paginilor (ex: 3,1,2,5,4)")
        order = [int(x.strip()) - 1 for x in seq_str.split(",") if x.strip().isdigit()]
        order = [i for i in order if 0 <= i < total]
        if not order:
            raise ValueError("Secventa de reordonare invalida")
        for i in order:
            writer.add_page(reader.pages[i])

    elif action == "optimize":
        for page in reader.pages:
            writer.add_page(page)
        writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)

    elif action == "watermark":
        wm_text = kwargs.get("watermark_text", "WATERMARK")
        # Create watermark page using fpdf2
        from fpdf import FPDF
        wm_pdf = FPDF()
        wm_pdf.add_page()
        wm_pdf.set_font("Helvetica", "B", 50)
        wm_pdf.set_text_color(200, 200, 200)
        wm_pdf.rotate(45, wm_pdf.w / 2, wm_pdf.h / 2)
        wm_pdf.text(wm_pdf.w / 4, wm_pdf.h / 2, wm_text)
        wm_bytes = wm_pdf.output()
        wm_reader = PdfReader(io.BytesIO(wm_bytes))
        wm_page = wm_reader.pages[0]
        for page in reader.pages:
            page.merge_page(wm_page)
            writer.add_page(page)

    else:
        raise ValueError(f"Operatie PDF necunoscuta: {action}")

    buf = io.BytesIO()
    writer.write(buf)
    return {"data": buf.getvalue(), "mime": "application/pdf", "filename": f"{_stem(name)}_{action}.pdf"}


# --- Router ---

def process(files: list[dict], operation: str, target_format: str, **kwargs) -> dict:
    if not files:
        raise ValueError("Nu au fost trimise fisiere")

    f0 = files[0]
    name = f0.get("name", "file.bin")
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    data = f0["data"]

    if operation == "merge":
        return merge_pdfs(files)
    if operation == "split":
        return split_pdf(data, kwargs.get("page_range", ""))
    if operation == "compress":
        return compress_pdf(data, name)
    if operation == "edit-pdf":
        return edit_pdf(data, name, **kwargs)

    if operation != "convert" or not target_format:
        raise ValueError(f"Operatiune invalida: {operation}")

    # Image conversions
    if ext in ("jpg", "jpeg", "png") and target_format == "pdf":
        return image_to_pdf(data, name)
    if ext in ("jpg", "jpeg", "png") and target_format in ("jpg", "jpeg", "png"):
        return image_convert(data, name, target_format)

    routes = {
        ("pdf", "docx"): lambda: pdf_to_docx(data, name),
        ("pdf", "html"): lambda: pdf_to_html(data, name),
        ("docx", "html"): lambda: docx_to_html(data, name),
        ("docx", "pdf"): lambda: docx_to_pdf(data, name),
        ("md", "html"): lambda: md_to_html(data, name),
        ("md", "pdf"): lambda: md_to_pdf(data, name),
        ("html", "pdf"): lambda: html_to_pdf(data, name),
        ("html", "md"): lambda: html_to_md(data, name),
        ("html", "docx"): lambda: html_to_docx(data, name),
    }
    key = (ext, target_format)
    if key in routes:
        return routes[key]()
    raise ValueError(f"Conversie {ext} -> {target_format} nu e suportata")


# --- Multipart parser ---

def parse_boundary(content_type: str) -> str:
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            return part[len("boundary="):].strip('"').strip("'")
    raise ValueError("No boundary in Content-Type")


def parse_multipart(body: bytes, boundary: str) -> dict:
    result: dict = {"files": [], "operation": "convert", "target_format": "", "page_range": ""}
    boundary_bytes = f"--{boundary}".encode()
    sections = body.split(boundary_bytes)

    for section in sections[1:]:
        if section.strip() in (b"--", b"", b"--\r\n"):
            continue
        header_end = section.find(b"\r\n\r\n")
        if header_end == -1:
            continue
        header = section[:header_end].decode("utf-8", errors="replace")
        content = section[header_end + 4:]
        if content.endswith(b"\r\n"):
            content = content[:-2]

        name_match = re.search(r'name="([^"]+)"', header)
        if not name_match:
            continue
        name = name_match.group(1)

        if name in ("operation", "target_format", "page_range",
                     "pdf_action", "rotate_angle", "watermark_text", "reorder_sequence"):
            result[name] = content.decode("utf-8").strip()
        elif name == "files" or "filename" in header:
            fname_match = re.search(r'filename="([^"]*)"', header)
            fname = fname_match.group(1) if fname_match else "file.bin"
            ct_match = re.search(r"Content-Type:\s*(\S+)", header)
            mime = ct_match.group(1) if ct_match else "application/octet-stream"
            result["files"].append({"name": fname, "data": content, "mime": mime})

    return result


# --- Handler ---

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", os.environ.get("ALLOWED_ORIGIN", "https://traduceri-matematica-7sh7.onrender.com"))
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        try:
            content_type = self.headers.get("Content-Type", "")
            content_length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(content_length)

            if "multipart/form-data" in content_type:
                boundary = parse_boundary(content_type)
                parts = parse_multipart(body, boundary)
            else:
                parts = json.loads(body)

            files = parts.get("files", [])
            operation = parts.get("operation", "convert")
            target_format = parts.get("target_format", "")
            extra = {k: parts[k] for k in ("page_range", "pdf_action", "rotate_angle",
                                             "watermark_text", "reorder_sequence")
                     if parts.get(k)}

            print(f"[CONVERT] {operation}: {len(files)} files, target={target_format}", file=sys.stderr)
            _log_to_file(f"ACTION  | Conversie initiata | {operation} | {len(files)} fisier(e) | Target: {target_format}")

            result = process(files, operation, target_format, **extra)

            out_data = result["data"]
            _log_to_file(f"OK      | Conversie completa | {result['filename']} | {len(out_data)} bytes")
            _log_to_file("")
            print(f"[CONVERT OK] {operation}: {result['filename']} ({result['mime']}, {len(out_data)} bytes)", file=sys.stderr)

            self.send_response(200)
            self.send_header("Content-Type", result["mime"])
            self.send_header("Content-Disposition", f'attachment; filename="{result["filename"]}"')
            self.send_header("Content-Length", str(len(out_data)))
            self.send_header("Access-Control-Allow-Origin", os.environ.get("ALLOWED_ORIGIN", "https://traduceri-matematica-7sh7.onrender.com"))
            self.send_header("Access-Control-Expose-Headers", "Content-Disposition, Content-Length")
            self.end_headers()
            self.wfile.write(out_data)

        except Exception as e:
            error_msg = str(e)
            _log_to_file(f"ERROR   | Conversie esuata | {type(e).__name__}: {error_msg}")
            _log_to_file("")
            print(f"[CONVERT ERROR] {type(e).__name__}: {error_msg}", file=sys.stderr)
            traceback.print_exc(file=sys.stderr)
            error_body = json.dumps({"error": error_msg, "status": "error"}).encode()
            self.send_response(400 if isinstance(e, ValueError) else 500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", os.environ.get("ALLOWED_ORIGIN", "https://traduceri-matematica-7sh7.onrender.com"))
            self.send_header("Access-Control-Expose-Headers", "Content-Disposition")
            self.end_headers()
            self.wfile.write(error_body)
