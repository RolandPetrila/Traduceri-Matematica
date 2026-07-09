"""Regression tests for DOCX export figure embedding (convert.html_to_docx).

Guards against the silent R-EXPORT/R-MATH bug where the tag-strip deleted every
<img> data-URI figure, producing a DOCX with text but zero figures. Deterministic:
a tiny inline PNG, no network, no API keys.
"""

import base64
import io
import zipfile

import pytest

from convert import html_to_docx
from docx import Document
from PIL import Image


def _png_data_uri_b64(w=120, h=80, color="red"):
    """Real PNG the DOCX writer accepts (a degenerate 1x1 is rejected)."""
    buf = io.BytesIO()
    Image.new("RGB", (w, h), color).save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("ascii")


_PNG_1x1 = _png_data_uri_b64()


def _docx(html: str):
    res = html_to_docx(html.encode("utf-8"), "t")
    return Document(io.BytesIO(res["data"])), res["data"]


def test_figures_embedded_as_pictures():
    # Two DISTINCT crops (docx de-duplicates identical blobs into one media file).
    fig_red = _png_data_uri_b64(color="red")
    fig_blue = _png_data_uri_b64(color="blue")
    html = (
        "<h2>Titlu</h2>"
        f'<div class="figure"><img src="data:image/png;base64,{fig_red}" alt="fig1"></div>'
        "<p>Un paragraf.</p>"
        f'<div class="figure"><img src="data:image/png;base64,{fig_blue}" alt="fig2"></div>'
    )
    doc, data = _docx(html)
    # Both figures must survive as real embedded pictures.
    assert len(doc.inline_shapes) == 2
    media = [n for n in zipfile.ZipFile(io.BytesIO(data)).namelist() if n.startswith("word/media/")]
    assert len(media) == 2
    # Text must survive too.
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Titlu" in text and "Un paragraf." in text
    # Regression: closing-tag names must NOT leak as literal paragraphs (the
    # capturing-group re.split bug produced stray 'div'/'p'/'h2' body text).
    para_texts = [p.text.strip() for p in doc.paragraphs]
    assert not any(t in ("div", "p", "h1", "h2", "h3", "li", "tr") for t in para_texts)


def test_text_only_still_works():
    doc, _ = _docx("<h1>H</h1><p>doar text</p>")
    assert len(doc.inline_shapes) == 0
    assert any("doar text" in p.text for p in doc.paragraphs)


def test_malformed_base64_is_skipped_not_fatal():
    # 'A' is a single base64 char → base64.b64decode raises binascii.Error, so this
    # actually exercises the b64decode try/except guard (unlike '@@@' which decodes
    # leniently). The bad figure must be skipped and the rest of the doc survive.
    import binascii

    with pytest.raises(binascii.Error):
        base64.b64decode("A")  # sanity: this input genuinely raises
    html = '<div><img src="data:image/png;base64,A"></div><p>text ok</p>'
    doc, _ = _docx(html)
    assert any("text ok" in p.text for p in doc.paragraphs)
    assert len(doc.inline_shapes) == 0  # the undecodable figure was skipped


def test_svg_figure_gets_placeholder_not_silent_loss():
    # Inline SVG can't be embedded in DOCX, but it must not vanish silently (R-EXPORT).
    html = '<div><svg viewBox="0 0 10 10"><path d="M0 0 L10 10"/></svg></div><p>dupa</p>'
    doc, _ = _docx(html)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Figur" in text  # placeholder present
    assert "dupa" in text
